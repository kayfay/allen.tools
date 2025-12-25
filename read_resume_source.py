from docx import Document

def read_resume_content(path):
    try:
        doc = Document(path)
        print("--- PARAGRAPHS ---")
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text.strip())
        
        print("\n--- TABLES ---")
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    print(" | ".join(row_text))
                    
    except Exception as e:
        print(f"Error reading file: {e}")

if __name__ == "__main__":
    read_resume_content('Allen Tools 4 Resume.docx')
