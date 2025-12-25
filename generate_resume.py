import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set narrow margins for 1-page layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('---'):
            continue
            
        # Header (Name)
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 50, 100)
            
        # Subheader (Title)
        elif line.startswith('## '):
            # Check if this is the job title at the top or a section header
            # Section headers usually follow '---' or are standard sections like 'Professional Summary'
            text = line[2:].strip()
            
            # If it's the Job Title at the very top (usually 2nd line)
            if "Data Scientist" in text and "Specialist" in text and len(doc.paragraphs) < 5:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(80, 80, 80)
            else:
                # Standard Section Header
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(text.upper())
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 50, 100)
                # Add bottom border simulation (underline)
                run.underline = False # simple underline might look cheap, let's stick to color

        # Contact Info Line
        elif line.startswith('**Email'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            
            # Simple parsing of the contact line
            # Remove MD bold syntax
            clean_line = line.replace('**', '').replace('|', '  |  ')
            run = p.add_run(clean_line)
            run.font.size = Pt(10)

        # Job Titles (### )
        elif line.startswith('### '):
            # Format: **Title** | Company | Date
            text = line[4:].strip()
            # Remove bold markers
            text = text.replace('**', '')
            
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            
            run = p.add_run(text)
            run.bold = True

        # List Items
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:].strip()
            
            # Handle bold inside list items
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
        # Education Entries (No markdown headers, just bold text lines in our file)
        elif line.startswith('**Master') or line.startswith('**BS'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            clean_line = line.replace('**', '')
            p.add_run(clean_line)

        # Other Text (Summary)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(docx_path)
    print(f"Resume saved to {docx_path}")

if __name__ == "__main__":
    create_resume_docx('resume.md', 'Allen_Tools_Resume.docx')
